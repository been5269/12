from pathlib import Path


def main() -> int:
    docx_path = Path(r"c:\Users\been5\OneDrive\바탕 화면\영어기능지도_개발계획서(5월6일 수정).docx")
    out_path = Path(r"c:\Users\been5\OneDrive\바탕 화면\앱\dev_plan_extracted.txt")
    log_path = Path(r"c:\Users\been5\OneDrive\바탕 화면\앱\dev_plan_extract_log.txt")

    try:
        from docx import Document  # type: ignore
    except Exception as e:  # noqa: BLE001
        log_path.write_text(f"python-docx import failed: {e!r}\n", encoding="utf-8")
        raise

    try:
        doc = Document(str(docx_path))
    except Exception as e:  # noqa: BLE001
        log_path.write_text(f"open docx failed: {e!r}\npath={docx_path}\n", encoding="utf-8")
        raise
    lines: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # tables
    for ti, table in enumerate(doc.tables):
        lines.append("")
        lines.append(f"[TABLE {ti + 1}]")
        for row in table.rows:
            cells = [(" ".join((c.text or "").split())).strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))

    out_path.write_text("\n".join(lines), encoding="utf-8")
    log_path.write_text(f"OK: wrote {out_path} with {len(lines)} lines\n", encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

